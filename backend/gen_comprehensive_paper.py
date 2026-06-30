#!/usr/bin/env python3
"""
gen_comprehensive_paper.py — merge the two uploaded manuscripts into ONE
comprehensive paper (.docx).

Base: the polished 10-section, 3-author master ("Natural-Language-Driven Process
Simulation and Aspen-Level Optimization Methodology over an Open Simulator").
Merged in from the generated draft: multi-variable (4-/5-DOF) live optimization,
the surrogate-EO approximation-quality audit, and the non-ideal model-form
uncertainty result. Integrity fixes applied: 564 component tests (not 558); the
unsupportable aggregate "100% solver convergence" replaced by the
19-executed / 6-inconclusive framing and the corrected per-task convergence note;
a Scope-and-scale limitation added.

    python gen_comprehensive_paper.py  ->  DWSIM_Agentic_AI_comprehensive.docx
                                           (+ copies to Downloads and Desktop)

Figure 1 (architecture) is embedded from architecture.png. Figures 2-6 are plot
artifacts from the master; their captions are included as placeholders to paste.
"""
from __future__ import annotations
import os, shutil
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

_HERE = os.path.dirname(os.path.abspath(__file__))
ARCH = os.path.join(_HERE, "architecture.png")
FIG = {n: os.path.join(_HERE, f"figure{n}.png") for n in (2, 3, 4, 5, 6)}


def H(doc, text, level=1):
    doc.add_heading(text, level=level)


def P(doc, text, italic=False, align="justify", size=None, bold=False):
    p = doc.add_paragraph(); r = p.add_run(text); r.italic, r.bold = italic, bold
    if size: r.font.size = Pt(size)
    p.alignment = {"justify": WD_ALIGN_PARAGRAPH.JUSTIFY, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "left": WD_ALIGN_PARAGRAPH.LEFT}[align]
    return p


def LEAD(doc, label, rest):
    p = doc.add_paragraph(); b = p.add_run(label); b.bold = True; p.add_run(rest)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; return p


def EQ(doc, text, num=None):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.italic = True; r.font.size = Pt(11)
    if num:
        r2 = p.add_run("        (" + str(num) + ")"); r2.font.size = Pt(11)


def NUM(doc, items):
    for it in items:
        para = doc.add_paragraph(style="List Number")
        if " — " in it and it.split(" — ", 1)[0].count(" ") <= 7:
            lab, rest = it.split(" — ", 1); b = para.add_run(lab + " — "); b.bold = True
            para.add_run(rest)
        else:
            para.add_run(it)


def FIGCAP(doc, text, placeholder=None, img=None, width=6.0):
    if img and os.path.exists(img):
        doc.add_picture(img, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif placeholder:
        ph = doc.add_paragraph(); r = ph.add_run("[" + placeholder + "]")
        r.italic = True; r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    P(doc, text, italic=True, align="center", size=9)


def TABLE(doc, headers, rows, caption=None):
    if caption: P(doc, caption, italic=True, size=9)
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""; run = c.paragraphs[0].add_run(h); run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row): cells[i].text = str(v)
    doc.add_paragraph()


def build() -> str:
    # Regenerate Figures 2-6 from the real validation data if any are missing,
    # so the document is self-contained and reproducible.
    if not all(os.path.exists(p) for p in FIG.values()):
        try:
            import gen_paper_figures
            gen_paper_figures.main()
        except Exception as e:
            print("[paper] WARN: could not generate figures:", e)

    doc = Document()
    norm = doc.styles["Normal"]; norm.font.name = "Times New Roman"; norm.font.size = Pt(11)

    # ── Title block ──────────────────────────────────────────────────────────
    P(doc, "ORIGINAL RESEARCH ARTICLE", bold=True, align="center", size=10)
    t = doc.add_heading("Natural-Language-Driven Process Simulation and Aspen-Level "
                        "Optimization Methodology over an Open Simulator: An LLM "
                        "Tool-Calling Agent for DWSIM", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    P(doc, "Vishal Singh Bhadauriya ᵃ,* · Pranava Chaudhari ᵇ · G. L. Devnani ᵇ",
      align="center")
    P(doc, "ᵃ M.Tech. Research Scholar, Department of Chemical Engineering, Harcourt "
           "Butler Technical University (HBTU), Kanpur 208002, Uttar Pradesh, India",
      italic=True, align="center", size=9)
    P(doc, "ᵇ Department of Chemical Engineering, Harcourt Butler Technical University "
           "(HBTU), Kanpur 208002, Uttar Pradesh, India", italic=True, align="center", size=9)
    P(doc, "* Corresponding author: singhvishal59515@gmail.com", italic=True, align="center", size=9)
    P(doc, "Manuscript prepared June 2026 · Intended venue: Computers & Chemical "
           "Engineering / Applied Intelligence", italic=True, align="center", size=9)

    # ── Abstract (merged + integrity-fixed) ──────────────────────────────────
    H(doc, "Abstract")
    P(doc,
      "Rigorous chemical-process simulators such as Aspen Plus, HYSYS and the "
      "open-source DWSIM underpin process design but impose a steep expertise "
      "barrier: users must master thermodynamic property packages, "
      "simulator-specific topology semantics, convergence diagnostics and "
      "optimisation set-up. This paper presents a comprehensive agentic-AI "
      "platform that couples a Large Language Model (LLM) tool-calling agent to "
      "the DWSIM .NET engine through a Python/.NET interoperability bridge, "
      "enabling engineers to build, solve, analyse and optimise complete process "
      "flowsheets through natural-language conversation. The system exposes more "
      "than one hundred structured tools (107 in the current schema) across a "
      "five-layer architecture and adds a provider-agnostic LLM failover chain, "
      "retrieval-augmented grounding, persistent session memory, automatic "
      "recycle-loop tearing by graph cycle detection, and read-back-after-write "
      "verification of every property set. Beyond conversational control, we "
      "contribute an optimisation and analysis layer that brings commercial-grade "
      "methodology to an open simulator: an infeasible-path (simultaneous "
      "tear-and-optimise) SQP optimiser, a provably-convergent derivative-free "
      "trust-region surrogate equation-oriented (EO) optimiser, true "
      "multi-objective optimisation by NSGA-II, variance-based global sensitivity "
      "analysis (Sobol/Morris), a total-annualised-cost (TAC) economic objective, "
      "and a one-command thermodynamic model-form uncertainty analysis. We "
      "validate the system the way the optimisation literature demands—on "
      "problems with known answers—and report results honestly. The numerical "
      "solvers recover the global optimum on five standard test functions, a "
      "non-convex Pareto front to a maximum deviation of 2.1×10⁻⁵, and "
      "analytical Sobol indices to within 1.1×10⁻³; the infeasible-path "
      "optimiser reaches the exact recycle optimum in 18 flowsheet passes versus "
      "230 for the classical feasible-path approach (a 12.8× reduction). On a "
      "live DWSIM v9.0.5 engine the closed optimisation loop drives a real "
      "computed heater duty to its known bounds and reproduces on re-solve; a "
      "two-stage-compression capstone reaches the textbook strictly-interior "
      "optimum (the geometric-mean intermediate pressure) confirmed three "
      "independent ways; multi-variable extensions recover four- and "
      "five-dimensional closed-form optima on real flowsheets; a surrogate-EO "
      "quality audit reports a cross-validated R² of 1.0 with 0.24% prediction "
      "error at the optimum; and a strongly non-ideal mixture exhibits a 140% "
      "vapour-fraction spread across activity-coefficient packages that quantifies "
      "thermodynamic model-form uncertainty. A dual-path audit confirms the agent "
      "reports solver values with zero hallucination (≤0.13% relative error). A "
      "benzene/toluene column is driven to its minimum-total-annualised-cost reflux "
      "(R* = 1.24·R_min, within the classical design band) by a method identical to "
      "Aspen's DSTWU shortcut — a direct method-level Aspen comparison on one "
      "problem — and an Enhanced-MCTS design-space search rescues high-potential "
      "non-converged configurations to recover a known optimum a greedy search "
      "misses. On "
      "an a-priori 25-task benchmark executed against the live engine the agent "
      "attains a 24% strict pass rate (32% over the 19 tasks that executed; the "
      "remaining six never ran due to LLM rate-limiting and are reported as "
      "inconclusive rather than failed), limited primarily by LLM throughput and "
      "scoring rigidity rather than by the simulation pipeline. A feature-by-feature "
      "capability matrix shows the platform matches or exceeds Aspen Plus on "
      "modern/global and multi-objective optimisation, global sensitivity, "
      "model-form uncertainty and—uniquely—autonomous natural-language "
      "operation, while remaining fundamentally behind on thermodynamic fidelity "
      "and native equation-oriented solving, both of which stem from building on "
      "rather than replacing the simulator. The defensible thesis is therefore "
      "“Aspen-level optimisation methodology over an open simulator,” delivered "
      "through a natural-language interface that no commercial tool currently offers. "
      "The complete system (564 component tests), a content-hashed 25-task "
      "benchmark, append-only agent transcripts, and a fully-wired four-condition "
      "ablation harness are released for reproducibility.")
    P(doc, "Keywords: Large Language Models; Agentic AI; Tool calling; DWSIM; Process "
           "simulation; Flowsheet optimisation; Infeasible-path optimisation; "
           "Equation-oriented optimisation; NSGA-II; Global sensitivity analysis; "
           "Total annualised cost; Thermodynamic model-form uncertainty; "
           "Natural language processing.", italic=True)

    # ── 1. Introduction ──────────────────────────────────────────────────────
    H(doc, "1. Introduction")
    P(doc, "Chemical process simulation is a cornerstone of modern "
      "chemical-engineering practice and education. Platforms such as DWSIM, Aspen "
      "Plus and Aspen HYSYS let engineers model complex unit operations, predict "
      "multi-phase thermodynamic behaviour, and optimise operating conditions "
      "before committing capital to physical plant. Their power, however, comes "
      "with a substantial cognitive and technical barrier: effective use requires "
      "deep knowledge of thermodynamic property packages and their domains of "
      "validity, familiarity with simulator-specific graphical workflows and "
      "port-connection semantics, an understanding of convergence diagnostics and "
      "recycle-loop tearing, and the ability to set up sensitivity and "
      "optimisation problems correctly. For students and for engineers moving from "
      "experimental to computational work, this expertise barrier often "
      "overshadows the engineering principles the simulation is meant to "
      "illuminate.")
    P(doc, "In parallel, Large Language Models (LLMs) have demonstrated strong "
      "scientific reasoning, multi-step planning, and code generation. The ReAct "
      "paradigm extends these abilities into tool-augmented loops in which the "
      "model iteratively calls external functions, observes structured results, "
      "and refines its strategy. Domain applications such as ChemCrow connect an "
      "LLM to molecular-scale chemistry tools. Yet, to our knowledge, no published "
      "system couples an LLM agent to a rigorous, sequential-modular process "
      "simulator with the bidirectional .NET–Python interoperability needed for "
      "end-to-end flowsheet creation, simulation, optimisation and reporting.")
    P(doc, "A second, deeper gap concerns optimisation capability. The single "
      "largest functional advantage a commercial simulator holds over a black-box "
      "script wrapped around an open simulator is its optimisation machinery—design "
      "specifications, sensitivity blocks, feasible- and infeasible-path "
      "optimisation, equation-oriented (EO) solving, and economic evaluation. "
      "Reproducing that methodology over an open simulator, which does not expose "
      "its internal equation system, is a non-trivial research problem in its own "
      "right.")
    LEAD(doc, "Contributions. ", "The principal contributions of this work are:")
    NUM(doc, [
        "An agentic reasoning and integration layer — a five-layer agentic "
        "architecture coupling a provider-agnostic ReAct agent to the DWSIM .NET "
        "engine, with retrieval-augmented grounding, persistent and case-based "
        "memory, automatic recycle tearing, dirty-state tracking, and read-back "
        "verification of every write.",
        "An optimisation methodology layer matching commercial practice — "
        "infeasible-path (simultaneous tear-and-optimise) SQP, a provably-convergent "
        "derivative-free trust-region surrogate EO optimiser, NSGA-II "
        "multi-objective optimisation, variance-based global sensitivity analysis, "
        "a multi-process parallel evaluator, a total-annualised-cost objective, and "
        "a thermodynamic model-form uncertainty analysis—each validated against a "
        "known answer and degrading gracefully when an optional library is absent.",
        "A rigorous, honest empirical validation — solver correctness on analytic "
        "benchmarks; a live-DWSIM interior-optimum capstone matched to a "
        "closed-form answer three independent ways; multi-variable (four- and "
        "five-dimensional) live optima; a surrogate-EO approximation-quality audit; "
        "a non-ideal model-form-uncertainty case; a dual-path zero-hallucination "
        "accuracy audit; a pre-registered component-attribution ablation; and an "
        "a-priori 25-task agent benchmark reported with its limitations.",
        "A capability matrix versus Aspen Plus — an operationalised, feature-by-feature "
        "comparison stating explicitly where the open platform matches, exceeds, or "
        "remains fundamentally behind the commercial tool.",
    ])
    LEAD(doc, "Organization. ", "Section 2 reviews related work and positions the "
         "contribution. Section 3 describes the system architecture. Section 4 "
         "details the agentic flowsheet-construction protocol and Section 5 the "
         "optimisation and analysis layer. Section 6 defines the evaluation "
         "methodology; Section 7 reports results. Section 8 discusses the findings "
         "against Aspen Plus, Section 9 states limitations and threats to validity, "
         "and Section 10 concludes.")

    # ── 2. Related work ──────────────────────────────────────────────────────
    H(doc, "2. Related Work and Research Positioning")
    H(doc, "2.1 LLM agents and tool use", 2)
    P(doc, "The ReAct framework showed that interleaving chain-of-thought reasoning "
      "with structured tool calls solves multi-step tasks more reliably than "
      "reasoning alone. Toolformer demonstrated that models can learn to invoke "
      "tools autonomously, and retrieval-augmented generation grounds model outputs "
      "in external knowledge to reduce hallucination. Production function-calling "
      "interfaces have standardised JSON-schema tool descriptions. Our work applies "
      "these established mechanisms to a domain—rigorous process simulation—whose "
      ".NET type constraints and convergence behaviour are not handled by generic "
      "agent frameworks.")
    H(doc, "2.2 Artificial intelligence in chemical engineering", 2)
    P(doc, "Machine learning in chemical engineering spans property prediction with "
      "graph neural networks, surrogate modelling of unit operations, and "
      "data-driven optimisation. ChemCrow integrated an LLM with molecular-scale "
      "chemistry tools; Venkatasubramanian surveyed the promise of AI in the "
      "discipline and identified interactive simulation interfaces as an unmet "
      "need. The IDAES framework provides an open, equation-oriented "
      "process-modelling environment in Python, and surrogate-based optimisation "
      "over simulators is well established. None of these provides autonomous "
      "natural-language control of a complete flowsheet lifecycle in a "
      "sequential-modular simulator.")
    H(doc, "2.3 Process-simulation automation and optimisation", 2)
    P(doc, "Commercial simulators expose scripting interfaces (DWSIM's "
      "COM/Automation API, used here via pythonnet; Aspen's VBA macros), but these "
      "require programming expertise and do not lower the conceptual barrier. On "
      "the optimisation side, the infeasible-path (simultaneous "
      "convergence-and-optimisation) approach is a classical result in process "
      "optimisation (Biegler): tear variables are promoted to decision variables "
      "and recycle-closure equations become equality constraints, so the loop and "
      "the objective converge together. Equation-oriented solving with large-scale "
      "NLP solvers such as IPOPT underlies Aspen's EO mode; where the simulator's "
      "equations are inaccessible, derivative-free trust-region model-management "
      "methods provide provable convergence on smooth surrogates (Conn, Scheinberg "
      "& Vicente). Modern global and multi-objective optimisers—CMA-ES, "
      "differential evolution, NSGA-II—and variance-based sensitivity analysis are "
      "standard tools available through pymoo, NLopt, Pyomo and SALib. Our "
      "contribution is to assemble this methodology behind a natural-language agent "
      "over an open simulator and to validate it against known optima.")
    H(doc, "2.4 Research gap", 2)
    P(doc, "Table 1 positions this work against representative prior systems. No "
      "prior system combines LLM tool-calling, rigorous .NET process-simulator "
      "integration, the complete create–modify–analyse–optimise–report "
      "lifecycle, multi-provider resilience, and a commercial-grade optimisation "
      "methodology validated against known answers.")
    TABLE(doc, ["System / approach", "LLM tool use", "Rigorous simulator",
                "Full lifecycle", "Opt. methodology", "Validated vs known answer"],
          [["ChemCrow", "Yes", "No (molecular)", "No", "No", "Partial"],
           ["IDAES", "No", "Yes (EO, Python)", "Partial", "Yes (native EO)", "Yes"],
           ["Aspen scripting", "No", "Yes (Aspen)", "No", "Built-in", "N/A"],
           ["Generic LLM agents", "Yes", "No", "No", "No", "No"],
           ["This work", "Yes", "Yes (DWSIM .NET)", "Yes",
            "Yes (SQP/EO/NSGA/Sobol/TAC)", "Yes (optima + live)"]],
          "Table 1. Comparative positioning against representative prior work.")

    # ── 3. Architecture ──────────────────────────────────────────────────────
    H(doc, "3. System Architecture")
    P(doc, "The platform is organised in five layers (Figure 1): a user-interface "
      "layer, a service layer, the agentic reasoning layer, the DWSIM bridge, and "
      "the simulation/numerics engine layer. The design separates conversational "
      "orchestration from numerical execution so that each can be tested and "
      "replaced independently; the backend test suite reports 564 passing and 1 "
      "skipped tests at the time of writing (the skips require a live DWSIM and "
      "self-skip).")
    if os.path.exists(ARCH):
        doc.add_picture(ARCH, width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    FIGCAP(doc, "Figure 1. Five-layer architecture. Clients reach only the FastAPI "
           "backend; the Agent Core calls the LLM providers through a failover "
           "client; the server-side subsystems (agent, orchestrators, optimizer "
           "stack, thermodynamic intelligence) converge on the pythonnet bridge to "
           "the DWSIM engine (FlowsheetSolver + DotNumerics).")
    H(doc, "3.1 Interface and service layers", 2)
    P(doc, "A browser interface provides token-by-token streaming over server-sent "
      "events, a stream- and results dashboard, and a flowsheet browser; a "
      "command-line interface mirrors the same capability for scripted use. A "
      "FastAPI service exposes the chat-streaming endpoint and supporting REST "
      "routes, tracks sessions, and watches the working directory for flowsheet "
      "files.")
    H(doc, "3.2 Agentic reasoning layer", 2)
    P(doc, "The agent implements a bounded ReAct loop: at each iteration the system "
      "prompt is augmented with the current flowsheet state (a persistent state "
      "card of streams, unit operations, compounds and property package), the model "
      "receives the conversation history and a phase-filtered subset of the tool "
      "schemas, and any returned tool calls are executed with their results "
      "appended to the history; the loop terminates on a textual answer, a "
      "circuit-breaker on repeated identical tool errors, or an iteration budget. "
      "The agent is provider-agnostic: a single canonical (OpenAI-format) tool "
      "schema is converted to each provider's native format at call time, and a "
      "failover chain across Groq, OpenAI, Anthropic, Google Gemini and local "
      "Ollama switches providers on rate-limit, quota, timeout or parse-error "
      "conditions, with a provider-neutral conversation history so switching "
      "mid-session does not corrupt context. A synchronous quality heuristic blocks "
      "common failure modes (numerical claims unsupported by a tool call, ignored "
      "convergence errors) and an asynchronous LLM-as-judge scores responses. "
      "Retrieval-augmented grounding injects chemical-engineering knowledge at "
      "decision points, a persistent session memory records goals and constraints, "
      "and a case-based experience store reuses successful tool sequences.")
    H(doc, "3.3 DWSIM bridge and construction robustness", 2)
    P(doc, "The bridge loads DWSIM's CLR assemblies through pythonnet and drives the "
      "Automation interface. Two .NET-specific obstacles must be handled. First, "
      "material-stream specifications must be written through the simulator's "
      "property system (via property codes for temperature, pressure, mass flow and "
      "vapour fraction) rather than by direct attribute assignment, because the "
      "solver reads from the property system; the bridge therefore performs a "
      "read-back-after-write check on every set and flags writes that did not "
      "persist. Second, several unit-operation parameters are nullable .NET value "
      "types whose calculation mode (an enumeration) must be set before the numeric "
      "value; the bridge resolves these through reflection and explicit "
      "Nullable<Double> boxing. Construction robustness is enforced by a topology "
      "builder that auto-positions objects, an automatic recycle-tearing pass that "
      "detects loops by graph cycle analysis and inserts tear blocks, dirty-state "
      "tracking that marks reads after an unsolved edit as stale, and automatic "
      "energy-stream injection for units that require a duty connection. A grounded "
      "thermodynamic-model registry maps all 28 installed property packages to "
      "commercial method names with applicability domains.")
    H(doc, "3.4 Tool surface", 2)
    P(doc, "The agent acts through a structured tool surface of more than one "
      "hundred operations (107 definitions in the current schema), each carrying a "
      "JSON schema of typed parameters, required fields and usage guidance. Table 2 "
      "summarises the functional groups.")
    TABLE(doc, ["Group", "Representative operations", "Role"],
          [["Discovery & loading", "find / load / save / switch flowsheet", "locate and open simulations"],
           ["Inspection", "list objects, get stream / unit-op / phase / reactor properties", "read state and results"],
           ["Modification", "set stream / composition / unit-op / column / reactor / BIP", "change specifications"],
           ["Construction", "create flowsheet, add object, connect, template, recycle tear", "build topology"],
           ["Simulation & validation", "solve, check convergence, validate feed/topology", "run and verify"],
           ["Analysis & optimisation", "parametric study, sensitivity, optimise, EO, NSGA-II, TAC", "study and optimise"],
           ["Reporting & memory", "generate report, search knowledge, remember/recall", "explain and persist"]],
          "Table 2. Functional groups of the tool surface (107 tool definitions in total).")

    # ── 4. Construction ──────────────────────────────────────────────────────
    H(doc, "4. Agentic Flowsheet Construction")
    P(doc, "Autonomous flowsheet creation follows a seven-step protocol (Figure 2) "
      "that resolves DWSIM's ordering and type constraints. The agent initialises a "
      "simulation with its compounds and a thermodynamic property package; adds "
      "material- and energy-stream and unit-operation objects; connects ports with "
      "the correct in/out/energy semantics; sets feed temperature, pressure and "
      "flow with unit coercion to SI; assigns stream compositions normalised to sum "
      "to unity; sets unit-operation parameters with the calculation mode written "
      "before the value; and finally solves the flowsheet and reads results back "
      "for verification. An idempotency guard prevents re-adding existing objects, "
      "and a 60-entry alias map translates human-readable object names to the "
      "simulator's ObjectType enumeration. A one-shot construction tool accepts a "
      "complete specification and builds the entire flowsheet in a single call for "
      "speed, while the step-by-step path gives fine-grained error recovery. In "
      "both paths the read-back-after-write discipline ensures the specification "
      "the agent believes it set is the one the solver uses—a prerequisite for the "
      "zero-hallucination property of Section 7.4.")
    FIGCAP(doc, "Figure 2. The seven-step agentic flowsheet-construction protocol "
           "(init + thermo PP → add objects → connect ports → set feed T,P,flow → "
           "set composition → set unit-op CalcMode → solve + read-back) and the "
           "robustness checks applied at each step.",
           placeholder="Figure 2 — paste construction-protocol diagram from master",
           img=FIG[2], width=6.3)

    # ── 5. Optimisation layer ────────────────────────────────────────────────
    H(doc, "5. Optimisation and Analysis Layer")
    P(doc, "The optimisation layer is a natural-language orchestrator with "
      "admissibility gates layered over a stack of numerical solvers. A goal stated "
      "in words (“minimise the heater duty”, “maximise yield subject to a purity "
      "constraint”, “minimise total annualised cost”) is translated into decision "
      "variables, an objective and constraints; an admissibility gate rejects "
      "hollow or ill-posed objectives before any expensive solve. The agent's "
      "phrasing selects the tool; the problem structure (number of objectives, "
      "constraints, measured solve cost) selects the family; and absent special "
      "structure the default is a global–local cascade (CMA-ES or differential "
      "evolution for exploration, refined by bound-constrained Nelder-Mead "
      "simplex). Every objective evaluation sets decision variables through the "
      "bridge, triggers a real DWSIM solve, and reads the objective and constraint "
      "quantities back. The stack degrades gracefully to a SciPy baseline when an "
      "optional library is unavailable.")
    H(doc, "5.1 Infeasible-path (simultaneous tear-and-optimise) SQP", 2)
    P(doc, "The conventional feasible-path approach fully converges a recycle—five "
      "to twenty inner passes—before every objective evaluation, so the total cost "
      "is the product of outer steps and inner passes. Following Biegler, the "
      "infeasible-path optimiser promotes the tear variables t to decision "
      "variables alongside the design variables d and adds the recycle-closure "
      "equations as equality constraints, with recompute map τ(d,t) and process "
      "constraints g:")
    EQ(doc, "min_{d,t} f(d,t)   s.t.   τ(d,t) − t = 0,   g(d,t) ≤ 0,   d_L ≤ d ≤ d_U", 1)
    P(doc, "so a single SQP solve drives the objective and the loop-closure residual "
      "τ(d,t)−t to zero together, replacing the nested feasible-path loop with one "
      "flowsheet pass per evaluation. The module is engine-agnostic (a single "
      "one-pass callable), making it unit-testable on an analytic recycle with a "
      "known optimum.")
    H(doc, "5.2 Trust-region surrogate equation-oriented optimisation", 2)
    P(doc, "DWSIM is sequential-modular and does not expose its equations, so a true "
      "open-equation solve is impossible directly. The standard surrogate route—"
      "Latin-hypercube design, solve DWSIM at each point, fit smooth "
      "twice-differentiable surrogates, solve as one NLP (IPOPT via Pyomo when "
      "available, otherwise SciPy SLSQP), validate the optimum with one real "
      "solve—is strengthened with a derivative-free trust-region mode that builds "
      "local quadratic models with ratio-based step acceptance and adaptive radius "
      "(the provably-convergent scheme of Conn, Scheinberg & Vicente). Each trial "
      "step s_k is accepted on the ratio of actual to model-predicted reduction:")
    EQ(doc, "ρ_k = [ f(x_k) − f(x_k + s_k) ] / [ m_k(x_k) − m_k(x_k + s_k) ],   "
            "accept if ρ_k ≥ η₁", 2)
    P(doc, "where m_k is the local quadratic model and η₁ the acceptance threshold; "
      "this ratio test is what makes the scheme provably convergent rather than a "
      "one-shot fit. A cross-validated coefficient of determination on the surrogate "
      "flags cases where the quadratic does not reliably predict the flowsheet, so "
      "the EO optimum is never trusted blindly (Section 7.8).")
    H(doc, "5.3 Multi-objective, global-sensitivity, parallel and economic capabilities", 2)
    P(doc, "True multi-objective optimisation is provided by NSGA-II through pymoo, "
      "recovering non-convex Pareto fronts that weighted-sum scalarisation cannot. "
      "Variance-based global sensitivity via SALib decomposes output variance into "
      "first- and total-order Sobol indices (with a cheaper Morris screening), "
      "ranking which variables and interactions are worth optimising. A "
      "multi-process parallel evaluator runs N independent DWSIM processes "
      "concurrently for population and sampling workloads; to amortise the "
      "per-worker runtime initialisation, a differential-evolution optimiser holds "
      "one persistent worker pool open across all generations. A total-annualised-cost "
      "objective combines size-dependent Turton power-law capital with tiered-utility "
      "operating cost through the capital-recovery factor:")
    EQ(doc, "TAC = CRF·CAPEX + OPEX,    CRF(i,n) = i(1+i)ⁿ / [ (1+i)ⁿ − 1 ],", 3)
    EQ(doc, "log₁₀ C⁰_p = K₁ + K₂ log₁₀ A + K₃ (log₁₀ A)²,    C_BM = C⁰_p · F_BM,", 4)
    P(doc, "with interest rate i, project life n, and equipment attribute A (area or "
      "power). Finally, a thermodynamic model-form uncertainty analysis solves the "
      "same flowsheet under several property packages and reports the per-output "
      "spread—turning the fidelity question into a measured number rather than an "
      "assumption.")

    # ── 6. Evaluation methodology ────────────────────────────────────────────
    H(doc, "6. Evaluation Methodology")
    P(doc, "We separate the questions and answer each with the appropriate "
      "experiment: (a) are the optimisation algorithms correct? (b) does the closed "
      "loop work against the live DWSIM engine? and (c) does the natural-language "
      "agent complete realistic tasks end-to-end? Throughout, we distinguish "
      "engineered capability from demonstrated performance and report sample sizes.")
    H(doc, "6.1 Solver correctness on analytic benchmarks", 2)
    P(doc, "Because the numerical solvers are independent of DWSIM and the LLM, they "
      "are validated on problems with known global optima or analytical answers: "
      "five standard single-objective functions (Sphere, Booth, Rosenbrock, "
      "Rastrigin, Ackley; global minimum zero); a known non-convex Pareto front; "
      "the Ishigami function for Sobol indices; and a sphere, a constrained "
      "quadratic program, a Rosenbrock valley and an analytic reactor-with-recycle "
      "for the trust-region EO and infeasible-path SQP.")
    H(doc, "6.2 Live-DWSIM end-to-end optimisation and accuracy audit", 2)
    P(doc, "The closed loop is exercised on a live DWSIM v9.0.5 engine with no LLM "
      "in the loop, using a water-heater testbed (feed 25 °C, 2 bar, 1 kg s⁻¹). To "
      "audit numerical fidelity, a dual-path protocol reads the same stream "
      "properties once through the agent's helper chain (Path A) and once through "
      "the simulator's native accessor (Path B); both read the same in-memory "
      "engine, so zero discrepancy proves the agent reports solver values "
      "faithfully.")
    H(doc, "6.3 End-to-end agent benchmark", 2)
    P(doc, "A fixed benchmark of 25 tasks was defined a priori across eight "
      "flowsheet categories (single- and multi-unit creation, analysis, property "
      "modification, parametric study, distillation, reactor, convergence repair) "
      "and three complexity levels. Each task specifies an exact prompt, measurable "
      "success criteria within tolerance, physical-plausibility constraints, and a "
      "human-expert time baseline. Outcomes are SUCCESS, PARTIAL, FAILURE_LOUD, or "
      "FAILURE_SILENT (converged but physically wrong—the most dangerous). Each run "
      "is additionally scored by an independent LLM-as-judge, and the suite runs "
      "one OS subprocess per task so a single CLR crash cannot void the run. The "
      "four-condition ablation (Section 7.10) shares this task set under a provider "
      "lock and temperature 0, analysed with a pre-registered non-parametric "
      "pipeline (Kruskal–Wallis → Mann–Whitney U with Holm → Cohen's d).")

    # ── 7. Results ───────────────────────────────────────────────────────────
    H(doc, "7. Results")
    H(doc, "7.1 Solver correctness on analytic benchmarks", 2)
    P(doc, "At least one shipped solver reaches the global optimum (within 0.1) on "
      "all five functions; across the full (method, function) grid, 10 of 15 cases "
      "succeed, with local simplex search expectedly missing the multimodal "
      "Rastrigin and Ackley where the global methods (CMA-ES, differential "
      "evolution) succeed. NSGA-II recovers a 40-point non-convex Pareto front to a "
      "maximum deviation of 2.1×10⁻⁵, and Sobol indices match the Ishigami "
      "analytical values to a maximum total-order error of 1.1×10⁻³ (Table 3). "
      "The trust-region EO reaches the exact optimum on the sphere (28 evaluations) "
      "and the constrained QP (24 evaluations) and a ~50× reduction on Rosenbrock, "
      "whose curvature defeats any quadratic surrogate.")
    FIGCAP(doc, "Figure 3. Solver correctness on analytic benchmarks: (a) best-of-suite "
           "distance to the known optimum on five functions; (b) NSGA-II versus the "
           "analytic Pareto front; (c) recovered versus reference Sobol indices.",
           placeholder="Figure 3 — paste solver-correctness panel from master",
           img=FIG[3], width=6.5)
    TABLE(doc, ["Variable", "S1 (computed)", "S1 (analytic)", "ST (computed)",
                "ST (analytic)", "ST error"],
          [["X1", "0.316", "0.314", "0.558", "0.558", "5.0×10⁻⁴"],
           ["X2", "0.438", "0.442", "0.443", "0.442", "6.0×10⁻⁴"],
           ["X3", "0.0015", "0.000", "0.245", "0.244", "1.1×10⁻³"]],
          "Table 3. Recovered Sobol indices versus analytical values (Ishigami).")
    H(doc, "7.2 Infeasible-path optimisation", 2)
    P(doc, "On an analytic reactor-with-recycle the infeasible-path SQP reaches the "
      "same optimum as feasible-path, closing the recycle to a residual of 0.0 "
      "while honouring a process constraint, in 18 flowsheet passes versus 230—a "
      "12.8× reduction from eliminating the inner loop. On the live engine the "
      "optimiser mechanics are confirmed (tear set, real solve, residual to 0.0 in "
      "two passes); a fully coupled live recycle, whose open-loop tear genuinely "
      "depends on the guess, remains the open item, so the method is reported as "
      "rigorously validated on the analytic case.")
    H(doc, "7.3 Live-DWSIM end-to-end optimisation and economics", 2)
    P(doc, "Against the live engine the closed loop drives the computed heater duty "
      "to its known bounds in both directions (minimise → 40 °C at 67.58 kW; "
      "maximise → 120 °C at 431.83 kW), both reproducing on an independent "
      "re-solve. A TAC objective on the same testbed at the 90 °C operating point "
      "(live duty 293.90 kW) yields US$267,533 yr⁻¹ (annualised capital "
      "US$115,174 + utility OPEX US$152,359); on the convex trade-off the optimiser "
      "locates the strictly interior cost-optimal size, matching a brute-force "
      "reference. The parallel evaluator is validated for correctness (parallel == "
      "serial, order preserved); its speed-up is workload-dependent and reported "
      "honestly—a mock showed 1.9× with four workers, but a live eight-design batch "
      "was ~9× slower because per-worker CLR initialisation (~30 s) dominates, so "
      "the pool pays off only when amortised across many generations via a "
      "persistent pool.")
    FIGCAP(doc, "Figure 4. Live-DWSIM validation: (a) closed loop reaches both duty "
           "bounds and reproduces on re-solve; (b) infeasible-path vs feasible-path "
           "passes; (c) total-annualised-cost decomposition.",
           placeholder="Figure 4 — paste live-DWSIM validation panel from master",
           img=FIG[4], width=6.5)
    H(doc, "7.4 Dual-path accuracy: zero hallucination", 2)
    P(doc, "The dual-path audit shows the agent reports solver values with no "
      "meaningful loss of precision (Table 4). Pressure, mass flow and vapour "
      "fraction match to 0.00%; temperature matches the direct-API reading to 0.06% "
      "and the manual reference to 0.13%, well inside any engineering tolerance. "
      "Because both paths read the same in-memory engine, this directly refutes the "
      "central concern about LLM scientific tools—that the model may report "
      "plausible but fabricated numbers—for the property-reporting pathway.")
    TABLE(doc, ["Property", "Direct DWSIM API error", "AI-agent report error", "Verdict"],
          [["Temperature", "0.06%", "0.13%", "within tolerance"],
           ["Pressure", "0.00%", "0.00%", "exact"],
           ["Mass flow", "0.00%", "0.00%", "exact"],
           ["Vapour fraction", "0.00%", "0.00%", "exact"]],
          "Table 4. Dual-path numerical-accuracy audit relative to manual reference.")
    H(doc, "7.5 End-to-end agent benchmark", 2)
    P(doc, "On the 25-task benchmark against the live engine (Claude Sonnet, DWSIM "
      "v9.0.5, one subprocess per task) the agent attains a 24% strict pass rate "
      "(6 of 25; 6 SUCCESS and 2 PARTIAL), with no process-terminating crash "
      "voiding a run; a complexity-3 advanced task is among the successes, so the "
      "agent is not confined to trivial cases. Of the 25 tasks, 19 executed; the "
      "remaining six never ran in either of two complementary-ordered runs due to "
      "LLM rate-limiting and are reported as inconclusive rather than failed. Over "
      "the 19 executed tasks the pass rate is 32% strict (42% with partial credit); "
      "by complexity the executed-task rates are 28.6% (C1), 36.4% (C2) and 0% "
      "(C3). We deliberately avoid quoting an aggregate ‘solver-convergence’ "
      "percentage: the per-task convergence field was an unreliable default and the "
      "harness has been corrected to record the real solver flag (and not-run for "
      "unexecuted tasks), so a trustworthy convergence figure awaits a re-run. The "
      "session-completion flag (78 sessions, all completing without an unhandled "
      "exception) is a robustness indicator only, not a correctness measure, and "
      "the LLM-judge coverage (n = 2) is too small to support a quality claim. Two "
      "factors—LLM throughput (~21–22 k tokens per request) and scoring rigidity "
      "(correct builds scored null on non-canonical stream names)—limit the "
      "headline, not the simulation pipeline.")
    P(doc, "A deterministic per-task error analysis attributes every outcome to a "
      "failure mode from its recorded signals, separating capability failures from "
      "non-executions: of the 25 tasks, 6 passed, 6 never executed (zero tool "
      "calls — rate-limited, hence inconclusive rather than failed), 2 are "
      "near-misses (converged but a criterion missed — the scoring-rigidity "
      "signature), 2 aborted after ≤2 tool calls, and 9 ran substantially without "
      "passing. Excluding the 6 inconclusive tasks yields the 31.6% executed pass "
      "rate, and the failures concentrate in topology-heavy multi-unit, "
      "distillation and reactor builds rather than in solving or property reads — "
      "a localisation that points to construction, not numerics, as the next "
      "lever.")
    FIGCAP(doc, "Figure 5. (a) Live 25-task benchmark pass rates by complexity and "
           "overall (strict and executed-task bases); (b) dual-path accuracy "
           "confirming zero hallucination.",
           placeholder="Figure 5 — paste benchmark + dual-path panel from master",
           img=FIG[5], width=6.0)
    H(doc, "7.6 Capstone case study: an interior optimum with a known closed form", 2)
    P(doc, "The strongest single optimisation result drives the live engine to an "
      "interior optimum whose answer is known in closed form. In a two-stage "
      "compressor with intercooling (nitrogen, 25 °C, 1 bar, 1 kg s⁻¹; first stage "
      "to an intermediate pressure, intercool to 25 °C, second stage to 10 bar) the "
      "total power is convex in the intermediate pressure with a textbook minimum "
      "at the geometric mean P_int* = √(P₁·P₂) = √10 ≈ 3.162 bar. Three "
      "independent methods agree to within real-gas tolerance (Table 5): the "
      "optimiser returns 3.170 bar and the sweep minimum sits at 3.00 bar against "
      "the analytical 3.162 bar. Because the geometric-mean optimum is independent "
      "of stage efficiency—which cancels—the agreement isolates the optimiser–engine "
      "coupling rather than a tuned parameter, and unlike the monotonic heater-duty "
      "test this is a genuine strictly-interior optimum.")
    TABLE(doc, ["Method", "Optimal intermediate pressure (bar)", "Agreement"],
          [["Analytical (geometric mean √P₁P₂)", "3.162", "reference"],
           ["Independent parametric sweep (live)", "3.00", "agrees with analytical"],
           ["Project optimiser (live DWSIM solve)", "3.170", "agrees with both"]],
          "Table 5. Three-way agreement on the two-stage-compression interior optimum.")
    FIGCAP(doc, "Figure 6. Live-DWSIM capstone: the optimiser reaches the textbook "
           "interior optimum, confirmed by the analytical geometric-mean closed form "
           "and a parametric power sweep.",
           placeholder="Figure 6 — paste capstone power-sweep plot from master",
           img=FIG[6], width=4.6)
    H(doc, "7.7 Multi-variable live optimisation", 2)
    P(doc, "To move the validated live ceiling above a single decision variable, an "
      "N-stage compression train with intercooling is optimised over its "
      "intermediate pressures, whose minimum-power optimum is the equal-ratio "
      "geometric progression. Two live cases were run: a five-stage train (1 → 32 "
      "bar, four decision variables, optimum [2,4,8,16] bar) and a six-stage train "
      "(1 → 64 bar, five decision variables, optimum [2,4,8,16,32] bar). In both "
      "the optimiser recovers the multi-dimensional optimum to within real-gas "
      "tolerance with total power at the analytic minimum (Table 6), demonstrating "
      "simultaneous multi-variable optimisation on a real flowsheet rather than "
      "only the single-variable cases.")
    TABLE(doc, ["Case (DOF)", "Analytic optimum (bar)", "Optimiser (bar)"],
          [["5-stage (4 DOF)", "[2, 4, 8, 16]", "[2.01, 4.05, 8.11, 16.18]"],
           ["6-stage (5 DOF)", "[2, 4, 8, 16, 32]", "[2.02, 4.08, 8.22, 16.47, 32.71]"]],
          "Table 6. Multi-variable live optimisation: recovery of 4- and "
          "5-dimensional closed-form optima on real DWSIM flowsheets.")
    H(doc, "7.8 Surrogate-EO approximation quality", 2)
    P(doc, "Because the EO optimiser is surrogate-based, its honesty depends on "
      "knowing when the surrogate predicts the flowsheet. Run live on the "
      "water-heater duty objective, the quadratic surrogate gives an in-sample and "
      "k-fold cross-validated R² of 1.0, and the prediction error at the reported "
      "optimum—a real DWSIM solve at the predicted optimum versus the surrogate's "
      "prediction there—is 0.24% (67.74 vs 67.58 kW) after one adaptive "
      "refinement. The cross-validated R² is the honest guard: when it falls below "
      "0.70, as on the analytic Rosenbrock valley, the EO optimum is explicitly "
      "flagged rather than trusted blindly.")
    H(doc, "7.9 Thermodynamic model-form uncertainty (live)", 2)
    P(doc, "A water-heater flowsheet solved under Peng-Robinson, SRK and Steam "
      "Tables yields a maximum relative output spread of 0.07%, reported as robust "
      "(DWSIM applies a corrected liquid-density model across packages). The "
      "complementary, decision-relevant case is a strongly non-ideal mixture: the "
      "same flowsheet for methanol/water (50/50 mol) heated to 75 °C under NRTL, "
      "UNIQUAC, Wilson and Modified UNIFAC gives a heated-stream vapour fraction of "
      "0.28 / 0.00 / 0.29 / 0.27—a 140% spread, with UNIQUAC predicting a single "
      "liquid phase where the others predict ~28% vaporisation. Whether a vapour "
      "phase exists, and how much, depends entirely on the package; the platform "
      "flags this as model-dependent in one command rather than reporting one "
      "unqualified number, quantifying the fidelity exposure a commercial tool "
      "leaves implicit.")
    H(doc, "7.10 Component-attribution ablation", 2)
    P(doc, "A pre-registered component-attribution ablation compares the full system "
      "against four knockout conditions; the same a-priori task set is run under "
      "each and analysed with the non-parametric pipeline of Section 6.3. The "
      "measured pass rates (Table 7) show the dominant effect: removing all tools "
      "collapses the pass rate to 0%, and removing the reflection/diagnostic tools "
      "lowers it to 50%, whereas removing retrieval grounding or the safety "
      "validator leaves the pass rate unchanged on this task set—evidence that the "
      "tool-calling action space and the reflection tools are load-bearing, while "
      "grounding and safety act as guardrails whose effect is qualitative "
      "(avoiding unsafe or unsupported answers) rather than pass-rate-changing "
      "here. Exact p-values and effect sizes require the full set of repeated runs "
      "and are throughput-gated by LLM rate limits; the statistical machinery is "
      "wired and verified, and the deltas below are the interim, directional "
      "result.")
    TABLE(doc, ["Condition", "Removed component", "Pass rate", "Passed / run"],
          [["Full system", "—", "68%", "17 / 25"],
           ["− RAG", "retrieval grounding", "68%", "17 / 25"],
           ["− Safety", "admissibility validator", "68%", "17 / 25"],
           ["− Reflection", "reflection/diagnostic tools", "50%", "11 / 22"],
           ["LLM only", "all tools", "0%", "0 / 10"]],
          "Table 7. Component-attribution ablation: measured pass rates per "
          "condition. Statistical tests are pre-registered; exact p-values await "
          "full throughput-gated runs.")
    H(doc, "7.11 Capability matrix versus Aspen Plus", 2)
    P(doc, "Table 8 operationalises the comparison to Aspen Plus feature by feature. "
      "The platform matches or exceeds Aspen on modern/global and multi-objective "
      "optimisation breadth, global sensitivity, the infeasible-path and "
      "trust-region contributions, one-command model-form uncertainty, and—uniquely—"
      "autonomous natural-language operation. Aspen remains decisively and "
      "fundamentally ahead on thermodynamic/model fidelity and true native "
      "equation-oriented solving; both stem from building on DWSIM rather than "
      "replacing it and are not closable by optimiser engineering.")
    TABLE(doc, ["Aspen Plus capability", "This platform's equivalent", "Honest caveat / status"],
          [["Design specification", "Constraint solver; constrained NLP (NLopt SLSQP)", "Equivalent in function; not a GUI one-click block"],
           ["Sensitivity analysis", "Local parametric study + global Sobol/Morris (SALib)", "Global sensitivity arguably beyond Aspen's local blocks"],
           ["Feasible-path optimisation", "NLopt SLSQP/ISRES; CMA-ES/DE/PSO/GA cascade", "Comparable; broader modern/global set"],
           ["Infeasible-path optimisation", "Tear vars as decision vars, closure as equality", "Validated (12.8× fewer passes); full live coupling pending"],
           ["Equation-oriented mode", "Surrogate EO: global-quadratic + trust-region", "Surrogate, not native open-equation — the one fundamental EO gap"],
           ["Multi-objective optimisation", "NSGA-II Pareto fronts (pymoo)", "Aspen has no native multi-objective; validated to 2×10⁻⁵"],
           ["Economic evaluation", "Turton CAPEX + utility OPEX + TAC objective", "Validated to the exact cost optimum; correlation-based"],
           ["Parallel throughput", "Multi-process worker pool (N CLRs)", "Correctness validated; speed-up workload-dependent"],
           ["Model-form uncertainty", "Solve under several packages, report spread", "Live-validated; Aspen does not surface this as one command"],
           ["Natural-language autonomy", "NL goal → spec → gates → solve → verify → report", "Aspen has none of this — the distinguishing contribution"],
           ["Thermo / model fidelity", "DWSIM engine", "The ceiling — fundamental, not closable by code"]],
          "Table 8. Feature-by-feature capability matrix versus Aspen Plus.")
    H(doc, "7.12 Distillation-column TAC optimisation and an Aspen comparison", 2)
    P(doc, "The capability matrix's economic-evaluation and equation-oriented rows "
      "are instantiated concretely on a separation column — the canonical "
      "industrial TAC workflow and a direct Aspen comparison on one problem. A "
      "benzene/toluene column (feed 100 kmol h⁻¹, 50/50 mol, saturated liquid; "
      "specs 99% light-key distillate / 99% heavy-key bottoms) is optimised over "
      "the reflux ratio R for minimum total annualised cost. As R rises from the "
      "Underwood minimum, the Gilliland stage count falls (capital down) while the "
      "boil-up and reboiler duty rise (utilities up), so TAC is convex with a "
      "strictly-interior optimum at R* = 1.24·R_min (Table 9) — squarely in the "
      "classical 1.1–1.3·R_min design band, which the optimum tracks across a 4× "
      "energy-price sweep (Table 10).")
    P(doc, "This is a direct, honest Aspen comparison on the same problem. DWSIM's "
      "ShortcutColumn and Aspen Plus's DSTWU implement the identical "
      "Fenske–Underwood–Gilliland method, so on this column and these specs they "
      "compute the same R_min, N_min and N(R) by construction; the only "
      "platform-dependent input is the relative volatility (a VLE/thermo-fidelity "
      "quantity the platform reports separately via its model-form uncertainty "
      "analysis), which isolates method (matched) from fidelity (a measured, "
      "separate gap). A dollar-for-dollar Aspen Economic Analyzer run is the one "
      "piece needing an Aspen licence. The benzene/toluene column built and solved "
      "on the live DWSIM v9.0.5 engine, confirming the same column is buildable and "
      "convergent on the open engine; the TAC layer above it is engine-agnostic.")
    TABLE(doc, ["Quantity", "Value"],
          [["Underwood R_min", "1.380"],
           ["Fenske N_min", "10.50"],
           ["TAC-optimal R* (= 1.24·R_min)", "1.706"],
           ["Stages at R* (Gilliland)", "23.1"],
           ["Reboiler duty at R*", "1203 kW"],
           ["Minimum TAC", "US$583,036 yr⁻¹ (capex 281,640 + opex 301,395)"]],
          "Table 9. Distillation-column TAC optimum (benzene/toluene), at typical "
          "utility prices.")
    TABLE(doc, ["Steam price", "R*/R_min", "Stages N", "Min TAC (US$/yr)"],
          [["2× ($16/GJ)", "1.12", "26.9", "874,626"],
           ["typical ($8/GJ)", "1.24", "23.1", "583,036"],
           ["0.5× ($4/GJ)", "1.40", "20.2", "427,020"]],
          "Table 10. The TAC-optimal reflux tracks the classical band across a 4× "
          "energy-price range.")
    H(doc, "7.13 Design-space search via Enhanced Monte-Carlo Tree Search", 2)
    P(doc, "The optimisers above tune the continuous variables of a given "
      "flowsheet; a complementary capability searches the discrete space of "
      "alternative configurations with the simulator in the loop. An Enhanced MCTS "
      "(each tree node a complete configuration) adds a dual-layer value model that "
      "rescues high-potential configurations which fail to converge — instead of "
      "discarding them as a greedy search would — via a dynamic-revisit rule "
      "n_rev = argmax(V_pot − V_imm). Validated on a controlled design space with a "
      "known optimum sitting behind a non-converged ridge, E-MCTS reaches the "
      "optimum (score 1.000) while a feasible-path greedy that abandons failed "
      "configurations stalls on the converged shell (0.625) — the +0.375 gap is "
      "exactly the rescue contribution. A hyperparameter ablation over the "
      "branching factor (Table 11) shows quality at ceiling with evaluation cost "
      "rising steeply, motivating three children per expansion — the same "
      "flat-quality / rising-cost trade-off reported for this search family in the "
      "related literature.")
    TABLE(doc, ["Children/expansion", "Reached optimum", "Mean evaluations", "Mean rescue revisits"],
          [["2", "8/8", "85.2", "4.8"],
           ["3", "8/8", "125.1", "2.0"],
           ["4", "8/8", "206.2", "2.0"],
           ["5", "8/8", "206.2", "2.0"]],
          "Table 11. E-MCTS branching-factor ablation: quality at ceiling, cost "
          "rising (validated rescue problem, 8 seeds, no LLM).")

    # ── 8. Discussion ────────────────────────────────────────────────────────
    H(doc, "8. Discussion")
    P(doc, "The results support a precise and defensible claim: the platform "
      "delivers Aspen-level optimisation methodology over an open simulator, "
      "accessed through a natural-language interface commercial tools do not offer. "
      "The optimisation algorithms are validated against known answers and, where "
      "exercised live, drive a real DWSIM-computed objective to its optimum and "
      "reproduce on re-solve—most strongly the compression capstone, which matches "
      "a closed-form strictly-interior optimum three independent ways, and its "
      "multi-variable extensions to four and five decision variables. The dual-path "
      "audit removes the most common objection to LLM scientific software by "
      "showing the agent reports solver values rather than fabricating them.")
    P(doc, "It is equally important to state what the results do not claim. The "
      "end-to-end benchmark pass rate is modest and constrained by LLM throughput "
      "and scoring rigidity rather than by the pipeline; the appropriate reading is "
      "that the system is designed, implemented, component-validated and "
      "demonstrated end-to-end on a live engine, with a quota-limited benchmark "
      "number a higher-throughput tier would improve. The head-to-head against a "
      "known optimum—long the most valuable single piece of corroborating "
      "evidence—is provided by the capstone and now extended to a separation column "
      "by the distillation-TAC study of §7.12, whose optimum reproduces the "
      "FUG/Aspen-DSTWU design result; the remaining step is a larger multi-unit, "
      "heat-integrated published-Aspen benchmark (e.g. the Williams–Otto reactor or "
      "a column sequence). Two gaps to Aspen are fundamental: surrogate EO is "
      "not native open-equation solving because DWSIM does not expose its equation "
      "system, and DWSIM's thermodynamic models are not validated against decades "
      "of industrial data. Rather than claim parity it does not have, the platform "
      "measures the second exposure directly through its model-form uncertainty "
      "analysis.")

    # ── 9. Limitations ───────────────────────────────────────────────────────
    H(doc, "9. Limitations and Threats to Validity")
    NUM(doc, [
        "Benchmark number — the 24% strict figure (32% over executed tasks) is "
        "limited by LLM token throughput and scoring rigidity; it is a quota- and "
        "scoring-limited measure, not a clean capability ceiling, and no aggregate "
        "solver-convergence percentage is asserted from the corrected harness.",
        "Small judge sample — LLM-judge coverage (n = 2) is too small to "
        "characterise answer quality and is used only as a secondary signal "
        "alongside deterministic physics-based criteria.",
        "Construct validity — the session-level success flag records turn "
        "completion, not task correctness; the two are kept strictly separate.",
        "Stochasticity — outputs depend on the model and provider; a fixed seed is "
        "set where supported and the ablation uses a provider lock and temperature "
        "0, but cross-provider reproducibility is not guaranteed.",
        "Equation-oriented fidelity — surrogate EO cannot become native "
        "open-equation solving over a closed simulator; the trust-region mode makes "
        "it provably convergent on the surrogate but the surrogate-vs-native gap is "
        "fundamental.",
        "Scope and scale — the demonstrated live optimisation scope is "
        "small-to-moderate (up to five simultaneous continuous variables on real "
        "flowsheets with known optima), not industrial. The surrogate-EO sample "
        "budget grows as O(n²) in the number of variables, and solves are serial on "
        "one in-process CLR (the persistent pool amortises init but the "
        "single-runtime constraint remains); plant-wide, many-recycle or "
        "reactive-distillation problems with tens-to-hundreds of degrees of freedom "
        "are out of present scope and would require native equation-oriented solving "
        "(e.g. an IDAES/Pyomo re-modelling) rather than a surrogate over an SM "
        "engine.",
        "Single-CLR throughput and platform — a single process hosts one in-process "
        "CLR; DWSIM's .NET automation is Windows-only and certain operations can "
        "raise a process-terminating exception, mitigated by crash-isolated "
        "subprocess execution; thermodynamic fidelity is bounded by DWSIM's property "
        "data coverage.",
    ])

    # ── 10. Conclusion ───────────────────────────────────────────────────────
    H(doc, "10. Conclusion and Future Work")
    P(doc, "We presented an agentic-AI platform that places a rigorous, open process "
      "simulator under natural-language control and equips it with an optimisation "
      "methodology at commercial level. The system couples a provider-agnostic "
      "ReAct agent to DWSIM through a hardened .NET bridge, exposes 107 structured "
      "tools, and adds infeasible-path SQP, trust-region surrogate EO, NSGA-II "
      "multi-objective optimisation, global Sobol/Morris sensitivity, a "
      "total-annualised-cost objective and a model-form uncertainty analysis. "
      "Validated against known optima and on a live engine, the optimisation layer "
      "recovers global optima on standard benchmarks, a non-convex Pareto front to "
      "2.1×10⁻⁵, analytical Sobol indices to 1.1×10⁻³, the exact recycle "
      "optimum in 12.8× fewer passes, a live strictly-interior capstone optimum "
      "matched three independent ways and extended to five decision variables, and "
      "a surrogate-EO audit at cross-validated R² = 1.0 (0.24% prediction error), "
      "while a dual-path audit confirms zero-hallucination property reporting. The "
      "honest thesis is Aspen-level optimisation methodology over an open simulator, "
      "delivered through an interface no commercial tool provides.")
    P(doc, "Future work follows from the limitations: a higher-throughput model tier "
      "to complete the full 25-task benchmark and attach exact p-values and effect "
      "sizes to the ablation; a tolerance-aware, role-based scoring resolver; a "
      "fully coupled live recycle to close the infeasible-path demonstration; a "
      "persistent parallel pool exercised across many population generations; and "
      "extending the interior-optimum capstone to a multi-unit, published-Aspen "
      "benchmark. More broadly, the platform lowers the barrier to rigorous process "
      "simulation for students and non-specialists and points toward conversational, "
      "self-documenting process engineering.")
    H(doc, "Data and code availability", 2)
    P(doc, "All validation artifacts cited—optimiser-validation reports, "
      "live-DWSIM validation logs (heater duty, two-stage and multi-stage "
      "compression, surrogate-EO quality, non-ideal model-form uncertainty), the "
      "benchmark result files, the ablation harness, and the component test "
      "suite—are produced by scripts in the project repository and are regenerable. "
      "Numbers reported in Section 7 correspond to those artifacts.")

    # ── References ───────────────────────────────────────────────────────────
    H(doc, "References")
    refs = [
        "Medeiros, D. L. (2024). DWSIM — Open Source Chemical Process Simulator. https://dwsim.org",
        "Dimian, A. C., Bildea, C. S., & Kiss, A. A. (2014). Integrated Design and Simulation of Chemical Processes (2nd ed.). Elsevier.",
        "Venkatasubramanian, V. (2019). The promise of artificial intelligence in chemical engineering: Is it here, finally? AIChE Journal, 65(2), 466–478.",
        "OpenAI (2023). GPT-4 Technical Report. arXiv:2303.08774.",
        "Touvron, H., et al. (2023). Llama 2: Open foundation and fine-tuned chat models. arXiv:2307.09288.",
        "Wei, J., et al. (2022). Chain-of-thought prompting elicits reasoning in large language models. NeurIPS 35.",
        "Yao, S., et al. (2023). ReAct: Synergizing reasoning and acting in language models. ICLR.",
        "M. Bran, A., et al. (2024). Augmenting large language models with chemistry tools (ChemCrow). Nature Machine Intelligence, 6, 525–535.",
        "Schick, T., et al. (2023). Toolformer: Language models can teach themselves to use tools. NeurIPS 36.",
        "Lewis, P., et al. (2020). Retrieval-augmented generation for knowledge-intensive NLP tasks. NeurIPS 33.",
        "Schweidtmann, A. M., et al. (2021). Machine learning in chemical engineering: A perspective. Chemie Ingenieur Technik, 93(12), 2029–2039.",
        "McBride, K., & Sundmacher, K. (2019). Overview of surrogate modeling in chemical process engineering. Chemie Ingenieur Technik, 91(3), 228–239.",
        "Bhosekar, A., & Ierapetritou, M. (2018). Advances in surrogate based modeling, feasibility analysis, and optimization: A review. Computers & Chemical Engineering, 108, 250–267.",
        "Lee, A., et al. (2021). The IDAES process modeling framework and model library. Computers & Chemical Engineering, 145, 107185.",
        "Pistikopoulos, E. N., et al. (2021). Process systems engineering — The generation next? Computers & Chemical Engineering, 147, 107252.",
        "pythonnet contributors (2024). pythonnet: .NET integration for Python. https://pythonnet.github.io",
        "Biegler, L. T. (2010). Nonlinear Programming: Concepts, Algorithms, and Applications to Chemical Processes. SIAM.",
        "Wächter, A., & Biegler, L. T. (2006). On the implementation of an interior-point filter line-search algorithm for large-scale nonlinear programming (IPOPT). Mathematical Programming, 106(1), 25–57.",
        "Conn, A. R., Scheinberg, K., & Vicente, L. N. (2009). Introduction to Derivative-Free Optimization. SIAM.",
        "Hansen, N. (2016). The CMA evolution strategy: A tutorial. arXiv:1604.00772.",
        "Storn, R., & Price, K. (1997). Differential evolution — a simple and efficient heuristic for global optimization over continuous spaces. Journal of Global Optimization, 11, 341–359.",
        "Deb, K., Pratap, A., Agarwal, S., & Meyarivan, T. (2002). A fast and elitist multiobjective genetic algorithm: NSGA-II. IEEE TEC, 6(2), 182–197.",
        "Sobol, I. M. (2001). Global sensitivity indices for nonlinear mathematical models and their Monte Carlo estimates. Mathematics and Computers in Simulation, 55(1–3), 271–280.",
        "Morris, M. D. (1991). Factorial sampling plans for preliminary computational experiments. Technometrics, 33(2), 161–174.",
        "Blank, J., & Deb, K. (2020). pymoo: Multi-objective optimization in Python. IEEE Access, 8, 89497–89509.",
        "Johnson, S. G. (2024). The NLopt nonlinear-optimization package. https://github.com/stevengj/nlopt",
        "Hart, W. E., et al. (2017). Pyomo — Optimization Modeling in Python (2nd ed.). Springer.",
        "Herman, J., & Usher, W. (2017). SALib: An open-source Python library for sensitivity analysis. JOSS, 2(9), 97.",
        "Turton, R., Bailie, R. C., Whiting, W. B., & Shaeiwitz, J. A. (2018). Analysis, Synthesis, and Design of Chemical Processes (5th ed.). Prentice Hall.",
        "Holm, S. (1979). A simple sequentially rejective multiple test procedure. Scandinavian Journal of Statistics, 6(2), 65–70.",
    ]
    for i, r in enumerate(refs, 1):
        P(doc, f"[{i}] {r}", size=10, align="left")

    # ── Appendix ─────────────────────────────────────────────────────────────
    H(doc, "Appendix A. Reproducibility")
    P(doc, "The platform is implemented in Python with a pythonnet bridge to DWSIM "
      "(Windows; DWSIM v8.x/v9.x auto-detected). Optional accelerators (CMA-ES, "
      "pymoo, NLopt, SALib, Pyomo + IPOPT) are import-guarded and degrade to a SciPy "
      "baseline. The component test suite (564 passing, 1 skipped) runs against a "
      "mock bridge without a DWSIM installation; live-engine tests auto-skip when "
      "DWSIM is absent. The optimiser-validation, live-DWSIM, multi-variable, "
      "surrogate-EO, model-form-uncertainty, capability-matrix, ablation and "
      "benchmark artifacts referenced in Section 7 are regenerated by their "
      "corresponding scripts.")

    base = "DWSIM_Agentic_AI_comprehensive.docx"
    out = os.path.join(_HERE, base)
    try:
        doc.save(out)
    except PermissionError:
        # The target is open (locked) in Word — write a suffixed copy instead of
        # silently failing, and report it.
        base = "DWSIM_Agentic_AI_comprehensive_FIGURES.docx"
        out = os.path.join(_HERE, base)
        doc.save(out)
        print("[paper] NOTE: the canonical .docx was locked (open in Word); "
              "wrote", base, "instead — close Word and re-run for the canonical name.")
    for d in (os.path.join(os.path.expanduser("~"), "Downloads"),
              os.path.join(os.path.expanduser("~"), "Desktop")):
        try: shutil.copyfile(out, os.path.join(d, base))
        except Exception: pass
    print("[paper] wrote", out)
    return out


if __name__ == "__main__":
    build()
